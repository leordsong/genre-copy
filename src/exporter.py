"""导出模块 - 将图片和文案导出为 Word/Excel 文档"""

import io
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Font
from PIL import Image


def export_to_docx(
    images: list[str | Path] | None,
    content: str,
    output_path: str | Path | None = None,
    title: str = "朋友圈文案",
) -> Path:
    doc = Document()

    if content:
        content_para = doc.add_paragraph(content)
        for run in content_para.runs:
            run.font.size = Pt(12)
        doc.add_paragraph()

    if images:
        for i, img_path in enumerate(images, 1):
            try:
                img_path = Path(img_path)
                if img_path.exists():
                    with Image.open(img_path) as pil_img:
                        if pil_img.mode in ("RGBA", "P"):
                            pil_img = pil_img.convert("RGB")
                        img_byte_arr = io.BytesIO()
                        pil_img.save(img_byte_arr, format="JPEG", quality=85)
                        img_byte_arr.seek(0)

                        max_width = 4.0
                        max_height = 4.0
                        width, height = pil_img.size
                        ratio = min(max_width * 72 / width, max_height * 72 / height)
                        doc_width = width * ratio / 72

                        doc.add_picture(img_byte_arr, width=Inches(doc_width))
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[图片 {i} 加载失败: {e}]")

    if output_path is None:
        output_path = Path.cwd() / f"文案导出_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    else:
        output_path = Path(output_path)

    doc.save(output_path)
    return output_path


def export_to_excel(
    images: list[str | Path] | None,
    content: str,
    template_name: str = "文案",
    output_path: str | Path | None = None,
) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = template_name[:31]

    ws.column_dimensions["A"].width = 60

    ws["A1"].value = content or ""
    ws["A1"].alignment = Alignment(wrap_text=True, vertical="top")
    ws["A1"].font = Font(size=11)

    temp_files = []
    if images:
        ws.row_dimensions[1].height = 150
        for i, img_path in enumerate(images):
            try:
                img_path = Path(img_path)
                if img_path.exists():
                    with Image.open(img_path) as pil_img:
                        if pil_img.mode in ("RGBA", "P"):
                            pil_img = pil_img.convert("RGB")
                        
                        max_size = (200, 200)
                        pil_img.thumbnail(max_size, Image.Resampling.LANCZOS)
                        
                        temp_path = Path.cwd() / f"_temp_img_{i}_{datetime.now().strftime('%H%M%S')}.jpg"
                        pil_img.save(temp_path, format="JPEG", quality=85)
                        temp_files.append(temp_path)
                        
                        xl_img = XLImage(str(temp_path))
                        col_letter = chr(ord("B") + i)
                        cell = f"{col_letter}1"
                        ws.column_dimensions[col_letter].width = 30
                        ws.add_image(xl_img, cell)
            except Exception:
                pass

    if output_path is None:
        output_path = Path.cwd() / f"文案导出_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    else:
        output_path = Path(output_path)

    wb.save(output_path)
    
    for temp_file in temp_files:
        try:
            temp_file.unlink()
        except Exception:
            pass
    
    return output_path
